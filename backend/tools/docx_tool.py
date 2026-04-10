import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import settings
from tools.logo_generator import generate_logo


def generate_docx(title: str, content: str, brand: dict, filename: str) -> str:
    os.makedirs(settings.output_dir, exist_ok=True)
    doc = Document()

    primary_color = _hex_to_rgb(brand.get("primary_color", "#000000"))
    font_name = brand.get("font", "Calibri")

    # Logo
    try:
        logo_path = generate_logo(brand)
        doc.add_picture(logo_path, width=Inches(2.5))
        doc.add_paragraph()
    except Exception:
        pass

    # Title
    heading = doc.add_heading(title, level=1)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(*primary_color)
    run.font.name = font_name
    run.font.size = Pt(22)

    # Horizontal rule under title using brand color
    _add_colored_border(heading, primary_color)

    # Body paragraphs
    for para in content.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        # Treat lines starting with # as subheadings
        if para.startswith("#"):
            sub = doc.add_heading(para.lstrip("#").strip(), level=2)
            for r in sub.runs:
                r.font.color.rgb = RGBColor(*primary_color)
                r.font.name = font_name
                r.font.bold = True
        elif len(para.split("\n")) == 1 and len(para) < 60 and not para.endswith("."):
            sub = doc.add_heading(para.strip(), level=2)
            for r in sub.runs:
                r.font.color.rgb = RGBColor(*primary_color)
                r.font.name = font_name
                r.font.bold = True
        else:
            p = doc.add_paragraph(para)
            for r in p.runs:
                r.font.name = font_name
                r.font.size = Pt(11)

    path = os.path.join(settings.output_dir, filename)
    doc.save(path)
    return path


def _add_colored_border(paragraph, rgb: tuple):
    """Add a bottom border in brand color to a paragraph."""
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _hex_to_rgb(hex_color: str) -> tuple:
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
